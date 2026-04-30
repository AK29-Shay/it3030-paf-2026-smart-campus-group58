#!/usr/bin/env python3
from docx import Document

# Read sample report
doc = Document('Sample_Final_Report_GroupID.docx')
print('=== SAMPLE REPORT STRUCTURE ===\n')

for i, para in enumerate(doc.paragraphs[:60]):
    if para.text.strip():
        level = para.style.name
        text = para.text[:100]
        print(f'{level:25s}: {text}')

print('\n=== TOTAL PARAGRAPHS ===')
print(f'Total paragraphs: {len(doc.paragraphs)}')

print('\n=== TABLE INFORMATION ===')
print(f'Total tables: {len(doc.tables)}')
for i, table in enumerate(doc.tables[:3]):
    print(f'\nTable {i+1} rows: {len(table.rows)}, cols: {len(table.columns)}')
    if table.rows:
        print(f'  First row: {[cell.text[:50] for cell in table.rows[0].cells[:3]]}')

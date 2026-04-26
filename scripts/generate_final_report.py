from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT = Path("docs/IT3030_PAF_Assignment_2026_Group58_Final_Report.docx")


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc, rows, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        shade(cell, "D9EAF7")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    return table


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def build_report():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        doc.styles[style_name].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("IT3030 - Programming Applications and Frameworks\n")
    run.bold = True
    run.font.size = Pt(16)
    run = title.add_run("PAF Assignment 2026 - Semester 1\n")
    run.bold = True
    run.font.size = Pt(15)
    run = title.add_run("Smart Campus Operations Hub\n")
    run.bold = True
    run.font.size = Pt(20)
    run = title.add_run("Group 58 Final Report")
    run.bold = True
    run.font.size = Pt(16)

    add_table(
        doc,
        [
            ("Course", "IT3030 - Programming Applications and Frameworks"),
            ("Faculty", "Faculty of Computing - SLIIT"),
            ("Scenario", "Smart Campus Operations Hub"),
            ("Repository", "https://github.com/AK29-Shay/it3030-paf-2026-smart-campus-group58"),
            ("Backend", "Spring Boot REST API"),
            ("Frontend", "React client web application"),
            ("Database", "MongoDB Atlas"),
            ("Report Date", "26 April 2026"),
        ],
        ["Item", "Details"],
    )

    doc.add_heading("Team Members", level=2)
    add_table(
        doc,
        [
            ("Dakshika M G N", "IT23813984", "Facilities catalogue and resource management"),
            ("Chamya N D", "IT23848184", "Booking workflow and QR check-in"),
            ("Chamoda M S", "IT23832480", "Incident ticketing and attachments"),
            ("Christopher K K", "IT23827530", "Authentication, user management, OAuth, notifications"),
        ],
        ["Member", "Student ID", "Main Contribution"],
    )
    doc.add_page_break()

    sections = [
        (
            "1. Abstract",
            [
                "Smart Campus Operations Hub is a production-inspired university operations platform developed for the IT3030 PAF Assignment 2026. The system provides facilities and asset catalogue management, resource booking workflows, maintenance and incident ticketing, role-based access, notifications, and OAuth/JWT authentication. The backend is implemented as a Spring Boot REST API and the frontend is implemented as a React web application. Data is persisted in MongoDB Atlas."
            ],
        ),
        (
            "2. Project Scope And Objectives",
            [
                "Provide a single web platform for bookable campus facilities, equipment, maintenance tickets, and operational notifications.",
                "Implement RESTful APIs with layered architecture, validation, meaningful error handling, and role-based security.",
                "Build a usable React frontend for students, admins, and technicians.",
                "Persist all operational data in MongoDB Atlas instead of in-memory collections.",
                "Provide Postman collections, automated CI checks, and viva-ready documentation.",
            ],
        ),
    ]

    for heading, items in sections:
        doc.add_heading(heading, level=1)
        if len(items) == 1:
            doc.add_paragraph(items[0])
        else:
            add_bullets(doc, items)

    doc.add_heading("3. Functional Requirements", level=1)
    add_table(
        doc,
        [
            ("Facilities and Assets Catalogue", "Create, read, update, delete, search, and filter resources with metadata such as type, category, capacity, location, availability, and status."),
            ("Booking Management", "Users request bookings; admins approve/reject; the system blocks overlapping active bookings; approved bookings support QR regeneration and check-in."),
            ("Maintenance Ticketing", "Users create tickets with priority/category; admins assign technicians; technicians update status; users and staff add comments/replies."),
            ("Notifications", "Users receive notifications for booking decisions, ticket assignments/status changes, comments, and role changes."),
            ("Authentication and Authorization", "Local JWT login, Google OAuth2 login, roles USER/ADMIN/TECHNICIAN, and protected API/frontend actions."),
        ],
        ["Module", "Requirement"],
    )

    doc.add_heading("4. Non-Functional Requirements", level=1)
    add_table(
        doc,
        [
            ("Security", "JWT authentication, Google OAuth2, role checks, ignored .env secrets, upload validation, and controlled CORS origins."),
            ("Usability", "Role-based dashboards, searchable catalogues, clear workflow buttons, validation messages, and notification panel."),
            ("Scalability", "MongoDB Atlas persistence and modular services allow independent growth of bookings, resources, tickets, and notifications."),
            ("Maintainability", "Layered controllers, services, repositories, DTOs, enums, and dedicated Postman/README documentation."),
            ("Reliability", "Conflict detection, structured error responses, automated tests, and CI validation."),
        ],
        ["Quality Attribute", "Implementation"],
    )

    doc.add_heading("5. System Architecture", level=1)
    architecture = doc.add_paragraph()
    architecture.add_run("React/Vite Client -> Axios/API Services -> Spring Security -> REST Controllers -> Service Layer -> MongoRepository -> MongoDB Atlas").bold = True
    add_bullets(
        doc,
        [
            "The React client manages screens, forms, protected routes, and API calls.",
            "Spring Security validates JWTs and OAuth2 sessions before controller execution.",
            "Controllers expose REST endpoints and delegate business rules to services.",
            "Services implement workflow rules such as booking conflicts, role changes, ticket assignment, and notifications.",
            "Repositories persist documents in MongoDB Atlas collections.",
        ],
    )

    doc.add_heading("6. Backend Architecture", level=1)
    add_table(
        doc,
        [
            ("controller", "REST endpoints for auth, users, resources, bookings, tickets, and notifications."),
            ("dto", "Request and response models used to keep API payloads separate from persistence entities."),
            ("entity", "MongoDB document models annotated with @Document and @Id."),
            ("repository", "Spring Data MongoRepository interfaces."),
            ("services", "Business logic for authentication, booking, ticketing, QR generation, resources, users, and notifications."),
            ("config", "JWT filter/util, OAuth success handler, CORS/security config, seed data, dotenv loading."),
            ("exception", "Global exception handler with consistent status and message responses."),
        ],
        ["Package", "Responsibility"],
    )

    doc.add_heading("7. Frontend Architecture", level=1)
    add_bullets(
        doc,
        [
            "React Router manages public routes, authenticated routes, and role-specific dashboards.",
            "AuthContext stores JWT/user state and supports local and Google login flows.",
            "Axios services attach bearer tokens and centralize API base behavior.",
            "Resource, Booking, Ticket, Notification, and Admin pages map directly to backend workflows.",
            "ProtectedRoute prevents unauthorized frontend navigation for role-specific pages.",
        ],
    )

    doc.add_heading("8. Database Design", level=1)
    add_table(
        doc,
        [
            ("users", "User accounts, role, provider, OAuth id, password hash, notification setting, technician specialty."),
            ("resources", "Bookable rooms, labs, equipment, capacity, location, availability, and active/out-of-service status."),
            ("bookings", "Resource booking requests, requester email, attendees, time range, workflow status, QR path, check-in time."),
            ("tickets", "Incident reports, description, priority, status, image paths, creator, assigned technician, timestamps."),
            ("ticket_comments", "Comments and replies with ownership and timestamps."),
            ("ticket_assignment_history", "Assignment audit trail with assigned-by, from technician, to technician, reason, and timestamp."),
            ("notifications", "Recipient, notification type, title/message, reference id/type, read state, created time."),
        ],
        ["MongoDB Collection", "Stored Data"],
    )
    doc.add_paragraph("MongoDB collections are created lazily when documents are inserted. The application seeds users and resources on startup when APP_SEED_ENABLED=true.")

    doc.add_heading("9. Main API Endpoints", level=1)
    add_table(
        doc,
        [
            ("Auth", "POST /api/auth/login", "Local login; returns JWT and user profile."),
            ("Auth", "POST /api/auth/signup", "Create local user account."),
            ("Auth", "GET /api/auth/me", "Return current authenticated user."),
            ("Users", "GET /api/users", "Admin list of users."),
            ("Users", "PUT /api/users/{id}/role", "Admin role update."),
            ("Resources", "GET /api/resources", "Public list of resources."),
            ("Resources", "POST /api/resources", "Admin create resource."),
            ("Resources", "PUT /api/resources/{id}", "Admin update resource."),
            ("Resources", "DELETE /api/resources/{id}", "Admin delete resource."),
            ("Bookings", "POST /api/bookings", "Create booking for authenticated user."),
            ("Bookings", "GET /api/bookings/my-bookings", "Current user bookings."),
            ("Bookings", "PUT /api/bookings/{id}/approve", "Admin approval."),
            ("Bookings", "PUT /api/bookings/{id}/reject", "Admin rejection with reason."),
            ("Bookings", "PUT /api/bookings/checkin/{id}", "QR check-in flow."),
            ("Tickets", "POST /api/tickets", "Create incident ticket."),
            ("Tickets", "PUT /api/tickets/{ticketId}/assign", "Admin technician assignment."),
            ("Tickets", "PUT /api/tickets/{ticketId}/status", "Admin/technician status update."),
            ("Tickets", "POST /api/tickets/{ticketId}/comments", "Add ticket comment."),
            ("Notifications", "GET /api/notifications", "List user notifications."),
            ("Notifications", "PATCH /api/notifications/{id}/read", "Mark notification as read."),
        ],
        ["Area", "Endpoint", "Purpose"],
    )

    doc.add_heading("10. Security Design", level=1)
    add_bullets(
        doc,
        [
            "JWTs are generated after local login and OAuth2 login, then sent in Authorization headers.",
            "Google OAuth2 uses backend redirect URI /login/oauth2/code/google and frontend callback /oauth2/redirect.",
            "ADMIN actions such as resource writes, user management, booking approval, and ticket assignment are protected.",
            "TECHNICIAN and ADMIN roles can update ticket status.",
            "CORS is restricted to configured local frontend origins.",
            "Ticket images are limited by count, size, and MIME type.",
        ],
    )

    doc.add_heading("11. Validation And Error Handling", level=1)
    add_table(
        doc,
        [
            ("400 Bad Request", "Invalid request payloads, malformed values, and general validation errors."),
            ("401 Unauthorized", "Missing/invalid JWT or invalid login credentials."),
            ("403 Forbidden", "Authenticated user does not have the required role."),
            ("404 Not Found", "Requested resource, user, ticket, notification, or booking does not exist."),
            ("409 Conflict", "Booking overlap, invalid booking workflow transition, or resource booking conflict."),
        ],
        ["Status", "Meaning"],
    )

    doc.add_heading("12. Testing And Quality Evidence", level=1)
    add_bullets(
        doc,
        [
            "JUnit/Mockito tests cover authentication, notification service behavior, user controller behavior, OAuth success handling, and application context loading.",
            "Postman collections verify login, resource CRUD, booking workflow, conflict detection, ticket workflow, comments, notifications, and role updates.",
            "GitHub Actions runs backend tests with MongoDB, frontend lint/build, and validates Postman collection/environment JSON.",
            "Manual UI demo validates role dashboards, form submission, API integration, QR check-in, and notifications.",
        ],
    )

    doc.add_heading("13. Innovation And Enhancements", level=1)
    add_bullets(
        doc,
        [
            "QR code generation and check-in for approved bookings.",
            "Technician assignment history for ticket auditability.",
            "SLA-style ticket timing fields for first response and resolution tracking.",
            "Notification preferences and notification panel endpoints.",
            "Postman automation that captures tokens and IDs across collections.",
            "MongoDB Atlas migration with startup seed data for demo readiness.",
        ],
    )

    doc.add_heading("14. Team Contribution And Version Control", level=1)
    doc.add_paragraph("The repository is hosted on GitHub and includes commit history across multiple branches and pull requests. The module ownership below is used for individual viva explanation.")
    add_table(
        doc,
        [
            ("Dakshika M G N - IT23813984", "Facilities catalogue", "Resource entity, repository, resource UI pages, images, filtering, admin resource management."),
            ("Chamya N D - IT23848184", "Bookings", "Booking forms, admin approval dashboard, calendar/list views, QR/check-in workflow."),
            ("Chamoda M S - IT23832480", "Tickets", "Ticket CRUD, attachments, technician assignment, ticket details, status updates, comments."),
            ("Christopher K K - IT23827530", "Auth and notifications", "JWT/OAuth flows, user roles, profile/settings, notifications, Postman/report hardening."),
        ],
        ["Member", "Module", "Evidence Of Work"],
    )

    doc.add_heading("15. AI Usage Disclosure", level=1)
    doc.add_paragraph(
        "AI assistance from Codex/ChatGPT was used for code, configuration, documentation, Postman artifact, and report preparation support. "
        "The generated output was reviewed, tested, adapted, and accepted by the project team before inclusion. The team remains responsible for understanding and explaining the submitted implementation during progress reviews and viva."
    )

    doc.add_heading("16. Setup And Demonstration Instructions", level=1)
    add_numbered(
        doc,
        [
            "Configure backend/smartcampus/.env with MongoDB Atlas, JWT, OAuth, CORS, and seed values.",
            "Run backend: cd backend/smartcampus; .\\mvnw.cmd spring-boot:run.",
            "Run frontend: cd frontend/smartcampus; npm install; npm run dev.",
            "Open http://localhost:5173 and log in with seeded accounts.",
            "Import Postman collections and local environment, then run collections in the documented order.",
        ],
    )

    doc.add_heading("17. Limitations And Future Improvements", level=1)
    add_bullets(
        doc,
        [
            "Email delivery for password reset and notifications can be integrated with a production email provider.",
            "More integration tests can be added around booking ownership and ticket workflow authorization.",
            "Deployment can be added using a cloud frontend and backend hosting platform.",
            "More dashboard analytics can be added for resource utilization and ticket SLA performance.",
        ],
    )

    doc.add_heading("18. Conclusion", level=1)
    doc.add_paragraph(
        "The Smart Campus Operations Hub satisfies the assignment requirement for a Spring Boot REST API and React client web application. "
        "It implements the required resource, booking, ticketing, notification, authentication, authorization, database persistence, testing, and documentation expectations. "
        "The system is ready for local demonstration, Postman verification, and individual viva explanation."
    )

    doc.add_heading("Appendix A - Evidence And References", level=1)
    add_bullets(
        doc,
        [
            "Assignment PDF: PAF_Assignment-2026.pdf",
            "Postman verification pack: postman/collections and postman/environments",
            "Viva guide: docs/viva-guidance.md",
            "Root setup guide: README.md",
            "Google OAuth authorized redirect URI: http://localhost:8080/login/oauth2/code/google",
        ],
    )

    screenshot_path = Path(r"C:\Users\akshayan\Pictures\Screenshots\Screenshot 2026-04-26 110936.png")
    if screenshot_path.exists():
        doc.add_paragraph("Figure: Google OAuth client configuration evidence.")
        try:
            doc.add_picture(str(screenshot_path), width=Inches(6.2))
        except Exception:
            doc.add_paragraph(f"OAuth screenshot path: {screenshot_path}")

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.space_after = Pt(6)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_report()
    print(OUTPUT)
